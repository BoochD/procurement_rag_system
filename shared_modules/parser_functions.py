from collections import defaultdict
import re
from typing import Dict, List

from docx import Document


SPACE_CHARS = [
    "\u00a0",
    "\u202f",
    "\u2009",
    "\u2002",
    "\u2003",
    "\u2004",
    "\u2005",
    "\u3000",
    "\ufeff",
    "\xa0",
]
QUOTES_MAP = {
    "\u00ab": '"',
    "\u00bb": '"',
    "\u201c": '"',
    "\u201d": '"',
    "\u201e": '"',
    "\u201f": '"',
    "\u2033": '"',
    "\u2018": "'",
    "\u2019": "'",
    "\u201a": "'",
    "\u201b": "'",
    "\u2032": "'",
    "В«": '"',
    "В»": '"',
    "вЂњ": '"',
    "вЂќ": '"',
    "вЂћ": '"',
    "вЂџ": '"',
    "вЂ™": "'",
    "вЂ": "'",
    "Р’В«": '"',
    "Р’В»": '"',
    "РІР‚Сљ": '"',
    "РІР‚Сњ": '"',
    "РІР‚С›": '"',
    "РІР‚Сџ": '"',
    "РІР‚в„ў": "'",
    "РІР‚В": "'",
}
DASH_CHARS = [
    "\u2010",
    "\u2011",
    "\u2012",
    "\u2013",
    "\u2014",
    "\u2015",
    "\u2212",
    "\ufe58",
    "\ufe63",
    "\uff0d",
]


def normalize_text(text: str) -> str:
    """
    Нормализует текст: чистит пробелы, кавычки, тире и переносы строк.
    """
    if text is None:
        return ""
    out = text
    out = out.replace("\n", "; ")
    for sp in SPACE_CHARS:
        out = out.replace(sp, " ")
    for k, v in QUOTES_MAP.items():
        out = out.replace(k, v)
    for dash in DASH_CHARS:
        out = out.replace(dash, "-")
    out = re.sub(r"[ \t\f\v]*\n[ \t\f\v]*", "\n", out)
    out = re.sub(r"[ \t]{2,}", " ", out)
    out = re.sub(r"\s*,\s*(?:,\s*)+", ", ", out)
    out = re.sub(r"\s*;\s*(?:;\s*)+", "; ", out)
    out = re.sub(r";+\s*$", "", out)
    return out.strip()


def dedupe_merged_cells(row):
    """
    Убирает соседние дубли ячеек в строке таблицы, которые появляются из-за merged cells.

    Пример:
        ["КТРУ 1", "КТРУ 1", "КТРУ 2"] -> ["КТРУ 1", "КТРУ 2"]
    """
    cleaned = []
    prev_text = None
    for cell in row.cells:
        text = cell.text.strip()
        if text != prev_text:
            cleaned.append(text)
        prev_text = text
    return cleaned


def get_row_cells(row) -> List[str]:
    """
    Возвращает все ячейки строки без удаления дублей, чтобы сохранить
    исходные индексы колонок в таблицах со сложным merged-header.
    """
    return [cell.text.strip() for cell in row.cells]


def row_looks_like_data(cells: List[str]) -> bool:
    """
    Пытается отличить строку данных от строки заголовка.

    Нужно для таблиц с многоуровневым header: верхние строки объединяются
    в один логический заголовок, а первая строка с данными уже не участвует
    в формировании `selected_indexes`.
    """
    normalized = [normalize_text(cell) for cell in cells if normalize_text(cell)]
    if not normalized:
        return False

    if re.fullmatch(r"\d+", normalized[0]):
        return True

    code_patterns = (
        r"\d+(?:\.\d+){2,}-\d+",
        r"\d{2}\.\d{2}\.\d{2,}",
    )
    if any(re.search(pattern, cell) for pattern in code_patterns for cell in normalized):
        return True

    numeric_like_cells = sum(bool(re.search(r"\d", cell)) for cell in normalized)
    return numeric_like_cells >= max(2, len(normalized) // 2)


def build_table_header(rows: List[List[str]], max_header_rows: int = 3) -> tuple[List[str], int]:
    """
    Собирает плоский заголовок таблицы из одной или нескольких верхних строк.

    Для обычной таблицы вернёт первую строку как есть.
    Для таблицы со сложным header склеит несколько строк по каждой колонке:
    "Требования ... | Значение характеристики".
    """
    if not rows:
        return [], 0

    header_rows = []
    for row in rows[:max_header_rows]:
        if header_rows and row_looks_like_data(row):
            break
        header_rows.append([normalize_text(cell) for cell in row])

    if not header_rows:
        header_rows = [[normalize_text(cell) for cell in rows[0]]]

    width = max(len(row) for row in header_rows)
    header = []
    for idx in range(width):
        parts = []
        for row in header_rows:
            if idx >= len(row):
                continue
            cell = row[idx]
            if cell and cell not in parts:
                parts.append(cell)
        header.append(" | ".join(parts))

    return header, len(header_rows)


def parse_okpd_entries(text: str):
    """
    Парсит строку с ОКПД2 в список словарей `{"okpd2": ..., "name": ...}`.

    Пример:
        "ОКПД2: 31.01.12 - Стулья; 31.01.13 - Столы"
        -> [{"okpd2": "31.01.12", "name": "Стулья"}, ...]
    """
    result = []
    for item in text.split(":")[1].split(";"):
        item = item.strip()
        if not item:
            continue
        item = normalize_text(item)
        code, name = item.split(" - ", 1)
        result.append({
            "okpd2": code.strip(),
            "name": name.strip(),
        })
    return result


def parse_ktry_entries(text: str):
    """
    Парсит строку с КТРУ в список словарей `{"ktru_code": ..., "name": ...}`.

    Пример:
        "КТРУ: 31.01.12.150-00000003 - Тумба офисная"
        -> [{"ktru_code": "31.01.12.150-00000003", "name": "Тумба офисная"}]
    """
    ktru_pattern = re.compile(r"\b\d{2}\.\d{2}\.\d{2}\.\d{3}-\d{8}\b")
    result = []
    for item in text.split(":", 1)[1].split(";"):
        item = item.strip()
        if not item:
            continue
        item = normalize_text(item)
        code_match = ktru_pattern.search(item)
        try:
            _, name = item.split(" - ", 1)
        except Exception:
            name = "У данного КТРУ не указано наименование в таблице"

        if code_match:
            code = code_match.group(0)
        else:
            try:
                code = item.split(" - ", 1)[0].strip()
            except Exception:
                code = item.strip()
        result.append({
            "ktru_code": code.strip(),
            "name": name.strip(),
        })
    return result


def _clean_keyword_dict(items_by_key: Dict[str, list[str]]) -> str:
    """
    Склеивает словарь найденных значений в чистую строку без дублей и лишних `;`.

    Пример:
        {"ОКПД": ["ОКПД2: 31.01.12 - Стул;;", "ОКПД2: 31.01.12 - Стул"]}
        -> "ОКПД2: 31.01.12 - Стул"
    """
    clean_items = []
    for key in items_by_key:
        for item in items_by_key[key]:
            item = item.strip()
            item = re.sub(r"\s*;\s*", "; ", item)
            item = re.sub(r"(;\s*){2,}", "; ", item)
            item = item.rstrip("; ").strip()
            if item:
                clean_items.append(item)

    return "\n".join(dict.fromkeys(clean_items))


def _extract_keyword_windows(text: str, keywords: list[str], window: int = 90) -> str:
    """
    Ищет в тексте фрагменты от ключевого слова и захватывает ещё `window` символов вправо.

    Пример:
        _extract_keyword_windows("блабла... КТРУ 31.01.12.150-00000003 тумба офисная", ["КТРУ"])
        -> "КТРУ 31.01.12.150-00000003 тумба офисная"
    """

    matches: list[str] = []
    for keyword in keywords:
        pattern = re.compile(rf"({re.escape(keyword)}[\s\S]{{0,{window}}})", re.IGNORECASE)
        for match in pattern.findall(text):
            clean_match = re.sub(r"\s+", " ", match).strip(" ;,\n\t")
            if clean_match:
                matches.append(clean_match)

    return "\n".join(dict.fromkeys(matches))

def extract_ktru_block(text: str, tail_chars: int = 30, fallback_chars: int = 150) -> str:
    start_match = re.search(r"КТРУ\s*:", text, flags=re.IGNORECASE)
    if not start_match:
        return ""

    start = start_match.start()
    fragment = text[start:]

    ktru_pattern = r"\d+(?:\.\d+){3}-\d+"
    matches = list(re.finditer(ktru_pattern, fragment))

    if not matches:
        return fragment[:fallback_chars].strip()

    last_match = matches[-1]
    end = last_match.end() + tail_chars

    return fragment[:end].strip()

class PlanParser:
    """
    #### Парсер для документа "Заявка в план-график", где данные обычно лежат в таблице.
    """

    def __init__(self, path: str):
        self.path = path
        self.doc = Document(path)

    def extract_table_kv_from_docx(self) -> List[str]:
        """
        Извлекает первую таблицу документа как список строк формата `"ключ: значение"`.

        Пример:
            ["ОКПД2: 31.01.12 - Стулья", "Количество: 10"]
        """
        kv = defaultdict(list)
        if not self.doc.tables:
            return []

        table = self.doc.tables[0]

        for row in table.rows:
            cells = []
            for c in row.cells:
                txt = normalize_text(c.text)
                if txt == "-":
                    txt = "отсутствует"
                cells.append(txt)

            if not cells[0]:
                cells = cells[1:]

            if len(cells) < 2:
                continue
            key = normalize_text(cells[0])
            val = ", ".join(dict.fromkeys(cells[1:]))
            if key:
                kv[key].append(val)

        plain_text_lines = [f"{k}: {', '.join(vals)}" for k, vals in kv.items()]
        return plain_text_lines

    def extract_clean_text(self, chunk_size: int = 50) -> str:
        """
        Возвращает весь текст документа из параграфов одной строкой с разделением через пустую строку.
        """
        full_text = []
        for para in self.doc.paragraphs:
            text = normalize_text(para.text)
            if text:
                full_text.append(text)
        full_text = "\n\n".join(full_text)

        return full_text


class DocumentParser:
    """
    #### Универсальный парсер Word-документов: умеет вытаскивать текст, таблицы и фрагменты по ключевым словам.
    """

    def __init__(self, path: str):
        self.path = path
        self.doc = Document(path)

    def extract_clean_text(self) -> str:
        """
        Возвращает весь текст документа из параграфов.
        """
        full_text = []
        for para in self.doc.paragraphs:
            text = normalize_text(para.text)
            if text:
                full_text.append(text)
        return "\n\n".join(full_text)

    def extract_table_data(self) -> str:
        """
        Извлекает все строки всех таблиц в виде `"первая ячейка: остальные ячейки"`.

        Пример:
            "ОКПД2: 31.01.12 - Стулья\nКоличество: 10"
        """
        tables_kv = []
        for table in self.doc.tables:
            for row in table.rows:
                cells = [normalize_text(c.text) for c in row.cells]
                if len(cells) >= 2:
                    key = cells[0]
                    val = ", ".join(cells[1:])
                    tables_kv.append(normalize_text(f"{key}: {val}"))
        return "\n".join(tables_kv)

    def table_to_markdown(self):
        """
        Преобразует все таблицы документа в markdown-таблицы.

        Пример:
            "| Поле | Значение |\n| --- | --- |\n| ОКПД2 | 31.01.12 |"
        """
        md_tables = []
        for table in self.doc.tables:
            rows = []
            for row in table.rows:
                rows.append(dedupe_merged_cells(row))

            md = ""
            for i, row in enumerate(rows):
                md += "| " + " | ".join(row) + " |\n"
                if i == 0:
                    md += "| " + " | ".join(["---"] * len(row)) + " |\n"
            md_tables.append(md)

        return "\n\n".join(md_tables)
    
    def extract_tables_characteristics(self, keywords: List[str]) -> dict:
        """
        Находит в таблицах колонки, чьи заголовки содержат ключевые слова, и возвращает словарь

        >>> Пример:
            {
                "22.11.11.000-00000007": {"Категория использования шины": "Обычная "...},
            }
        """
        extracted_rows = []
        keyword_lower = [kw.lower() for kw in keywords]
        result = {}
        codes = []
        for table in self.doc.tables:
            rows = [get_row_cells(row) for row in table.rows if len(dedupe_merged_cells(row))>1]
            # print(rows)
            if not rows:
                continue

            header, header_rows_count = build_table_header(rows)
            selected_indexes = [
                idx for idx, cell in enumerate(header)
                if any(kw in cell.lower() for kw in keyword_lower)
            ]

            first_col_is_number = False
            if header:
                first_header = header[0].lower()
                first_col_is_number = "№" in header[0] or "п/п" in first_header or "номер" in first_header

            if not selected_indexes:
                continue
            
            i=0
            for row in rows[header_rows_count:]:
                
                normalized_row = [normalize_text(cell) for cell in row]
                selected_cells = [
                    normalized_row[idx]
                    for idx in selected_indexes
                    if idx < len(normalized_row) and normalized_row[idx]
                ]
                # assert len(selected_cells) == len(keywords), f"Не нашёл все колонки {keywords}"
                # print("Нашёл колонки:", selected_cells)
                if len(selected_cells) == 1 or "Дополнительные характеристики" in selected_cells[1]:
                    continue

                if (
                    first_col_is_number
                    and selected_cells[0].strip().isdigit()
                ):
                    num  = selected_cells[0]
                    code = f"№{num}. " + selected_cells[1].split()[0]
                    name = selected_cells[2]
                    val  = selected_cells[3]
                else:
                    code = selected_cells[0].split()[-1]
                    name = selected_cells[1]
                    val  = selected_cells[2]
                codes.append(code)
                
                if any(selected_cells):
                    result.setdefault(code, {}).update({name: val})
        # print(codes)
        return result, set(codes)
    
    def extract_tables_columns(self, keywords: List[str]) -> str:
        """
        Находит в таблицах колонки, чьи заголовки содержат ключевые слова, и возвращает их построчно.

        Пример:
            "| КТРУ: 31.01.12.150-00000003 | ОКПД2: 31.01.12 |"
        """
        extracted_rows = []
        keyword_lower = [kw.lower() for kw in keywords]

        for table in self.doc.tables:
            rows = [get_row_cells(row) for row in table.rows if len(dedupe_merged_cells(row))>1]
            # print(rows)
            if not rows:
                continue

            header, header_rows_count = build_table_header(rows)
            selected_indexes = [
                idx for idx, cell in enumerate(header)
                if any(kw in cell.lower() for kw in keyword_lower)
            ]
            if not selected_indexes:
                continue
            
            i=0
            for row in rows[header_rows_count:]:
                
                normalized_row = [normalize_text(cell) for cell in row]
                selected_cells = [
                    f"{header[idx]}: {normalized_row[idx]}"
                    for idx in selected_indexes
                    if idx < len(normalized_row) and normalized_row[idx]
                ]
                # print(selected_cells)
                if any(selected_cells):
                    i+=1
                    extracted_rows.append("| " + " | ".join(selected_cells) + " |") #"| " + str(i) +

        return "\n".join(dict.fromkeys(extracted_rows))

    def extract_rows_region(self, keyword: str, left_range: int = 1, right_range: int = 1) -> str:
        """
        Ищет строки таблиц с ключевым словом и возвращает найденную ячейку вместе с соседними.

        Пример:
            "| Наименование | КТРУ | 31.01.12.150-00000003 |"
        """
        extracted_rows = []
        keyword_lower = keyword.lower()

        for table in self.doc.tables:
            for row in table.rows:
                cells = [normalize_text(cell) for cell in dedupe_merged_cells(row)]
                if not cells:
                    continue

                matching_indexes = [
                    idx for idx, cell in enumerate(cells)
                    if keyword_lower in cell.lower()
                ]
                if not matching_indexes:
                    continue

                used_indexes = set()
                selected_cells = []

                for idx in matching_indexes:
                    start = max(0, idx - left_range)
                    end = min(len(cells), idx + right_range + 1)

                    for cell_idx in range(start, end):
                        if cell_idx not in used_indexes:
                            selected_cells.append(cells[cell_idx])
                            used_indexes.add(cell_idx)

                extracted_rows.append("| " + " | ".join(selected_cells) + " |")

        return "\n".join(extracted_rows)

    def extract_table_cells_by_keyword(self, keywords: List[str]) -> Dict[str, List[str]]:
        """
        Ищет в таблицах ячейки, содержащие ключевые слова, и группирует найденное по каждому ключу.

        Пример:     
            {"ОКПД": ["ОКПД2: 31.01.12 - Стулья"]}
        """
        results = defaultdict(list)
        for table in self.doc.tables:
            for row in table.rows:
                for kw in keywords:
                    cells = [normalize_text(c.text) for c in row.cells if kw.lower() in normalize_text(c.text).lower()]
                    if cells:
                        results[kw].append(", ".join(cells))
        return dict(results)


def parce_contracters_onmck(ONMCK_path: str) -> Dict[str, List[str]]:
    """
    Достаёт из таблицы ОНМЦК цены поставщиков по каждому наименованию товара.

    Возвращает словарь формата:
        {
            "Стол письменный": ["17 170,00", "14 586,42", "14 442,00"],
            ...
        }

    Если одно и то же наименование встречается несколько раз, цены из всех строк
    добавляются в один список в порядке следования по документу.
    """
    doc = Document(ONMCK_path)
    parsed_prices: Dict[str, List[str]] = defaultdict(list)
    print("num tables:", len(doc.tables))
    for table in doc.tables:
        
        rows = [get_row_cells(row) for row in table.rows if len(dedupe_merged_cells(row)) > 1]
        if not rows:
            continue
        
        header, header_rows_count = build_table_header(rows)
        if not header:
            continue
        
        name_col_idx = next(
            (
                idx
                for idx, cell in enumerate(header)
                if "наименование" in cell.lower()
            ),
            None,
        )

        if name_col_idx is None:
            continue

        

        supplier_price_indexes = [
            idx
            for idx, cell in enumerate(header)
            if ("поставщик" in cell.lower() or "исполнитель " in cell.lower()) and "цена за" in cell.lower()
        ]
        print(header)
        print(supplier_price_indexes)
        if not supplier_price_indexes:
            continue

        for row in rows[header_rows_count:]:
            normalized_row = [normalize_text(cell) for cell in row]
            if name_col_idx >= len(normalized_row):
                continue

            item_name = normalized_row[name_col_idx]
            if not item_name or item_name.lower() == "итого:":
                continue

            prices = [
                normalized_row[idx]
                for idx in supplier_price_indexes
                if idx < len(normalized_row) and normalized_row[idx]
            ]
            if prices:
                parsed_prices[item_name].extend(prices)

    return dict(parsed_prices)

def parse_price(text: str) -> float:
    text = normalize_text(text)
    return float(text.replace(' ', '').replace(',', '.'))

def parse_contracters_onmck_by_row_number(ONMCK_path: str) -> Dict[str, List[str]]:
    """
    Парсер ОНМЦК для таблиц с многострочным заголовком.

    Начало строк с данными определяется по первому столбцу:
    если там номер позиции вида 1, 1., 01, то это уже данные.
    Все строки выше объединяются в заголовок по индексам колонок.
    """
    doc = Document(ONMCK_path)
    parsed_prices: Dict[str, List[str]] = defaultdict(list)

    for table in doc.tables:
        rows: List[List[str]] = []
        for row in table.rows:
            normalized_row = [normalize_text(cell.text) for cell in row.cells]
            if any(normalized_row):
                rows.append(normalized_row)

        if not rows:
            continue

        data_start_idx = next(
            (
                idx
                for idx, row in enumerate(rows)
                if row and row[0] and re.fullmatch(r"\d+\.?", row[0])
            ),
            None,
        )
        if data_start_idx is None or data_start_idx == 0:
            continue

        header_rows = rows[:data_start_idx]
        width = max(len(row) for row in rows)
        header: List[str] = []
        for col_idx in range(width):
            parts = []
            for row in header_rows:
                if col_idx >= len(row):
                    continue
                cell = row[col_idx]
                if cell and cell not in parts:
                    parts.append(cell)
            header.append(" | ".join(parts))
        # print(header)
        name_col_idx = next(
            (
                idx
                for idx, cell in enumerate(header)
                if "наименование товара" in cell.lower() or cell.lower() == "наименование"
            ),
            None,
        )
        if name_col_idx is None:
            continue

        supplier_price_indexes = [
            idx
            for idx, cell in enumerate(header)
            if ("поставщик" in cell.lower() or "исполнитель " in cell.lower()) and "цена за ед" in cell.lower()
        ]
        # print(supplier_price_indexes)
        if not supplier_price_indexes:
            continue

        for row in rows[data_start_idx:]:
            if name_col_idx >= len(row):
                continue

            row_number = row[0] if row else ""
            item_name = "№" + row_number + " " + normalize_text(row[name_col_idx])
            if not item_name or item_name.lower() == "итого:" or row_number.lower() == "итого:":
                continue

            prices = [row[idx] for idx in supplier_price_indexes if idx < len(row) and row[idx]]
            # print(prices)
            if prices:
                prices = [parse_price(price) for price in prices]
                parsed_prices[item_name].extend(prices)

    return dict(parsed_prices)
