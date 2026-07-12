from __future__ import annotations

from collections import Counter, defaultdict
from decimal import Decimal
from io import BytesIO
import re
from typing import Callable, Iterable

from docx import Document
from docx.shared import Inches, Pt, RGBColor

from summary_model.domain.models import (
    ExtractedValue,
    Finding,
    FindingSeverity,
    FindingStatus,
    ProcurementPackage,
)


DOCUMENT_ORDER = (
    "plan",
    "request",
    "explanatory_note",
    "ooz",
    "contract",
    "onmck",
)
DOCUMENT_LABELS = {
    "plan": "Заявка в план-график",
    "request": "Обращение о проведении закупки",
    "explanatory_note": "Пояснительная записка",
    "ooz": "Описание объекта закупки (ООЗ)",
    "contract": "Проект контракта",
    "onmck": "Обоснование НМЦК",
}
STATUS_MARKERS = {
    FindingStatus.FAILED: "ОШИБКА",
    FindingStatus.UNCERTAIN: "ТРЕБУЕТ ПРОВЕРКИ",
    FindingStatus.SKIPPED: "ПРОВЕРКА ПРОПУЩЕНА",
}
STATUS_COLORS = {
    FindingStatus.FAILED: RGBColor(0xDC, 0x35, 0x45),
    FindingStatus.UNCERTAIN: RGBColor(0xD9, 0x77, 0x06),
    FindingStatus.SKIPPED: RGBColor(0x6C, 0x75, 0x7D),
}
FALLBACK_SECTIONS = (
    ("package.", "Комплектность пакета"),
    ("registry.ktru.", "Проверка КТРУ"),
    ("registry.okpd2.", "Проверка ОКПД2 и ПП №1875"),
    ("items.", "Согласованность позиций документов"),
    ("items_consistency.", "Согласованность позиций документов"),
    ("delivery_and_finance.", "Поставка, сроки и финансовые условия"),
    ("legal_and_completeness.", "Условия контракта и правовые требования"),
    ("price.", "Расчёт НМЦК и цены поставщиков"),
)


def _raw(value: ExtractedValue | None):
    if value is None:
        return None
    return (
        value.normalized_value
        if value.normalized_value is not None
        else value.raw_value
    )


def _clean_decimal(value: Decimal) -> str:
    rendered = format(value, "f")
    if "." in rendered:
        rendered = rendered.rstrip("0").rstrip(".")
    return rendered or "0"


def _display_scalar(value) -> str:
    if value is None:
        return ""
    if isinstance(value, Decimal):
        return _clean_decimal(value)
    if isinstance(value, dict):
        if "amount" in value:
            amount = _display_scalar(value.get("amount"))
            return f"{amount} {value.get('currency') or ''}".strip()
        if "value" in value:
            amount = _display_scalar(value.get("value"))
            return f"{amount} {value.get('unit') or ''}".strip()
    return str(value).strip()


def _value_text(value: ExtractedValue | None) -> str:
    return _display_scalar(_raw(value))


def _documents(package: ProcurementPackage):
    for key in DOCUMENT_ORDER:
        document = getattr(package, key, None)
        if document is not None:
            yield key, document
    for index, document in enumerate(package.commercial_offers, start=1):
        yield f"commercial_offer_{index}", document


def _document_label(key: str, document) -> str:
    if key.startswith("commercial_offer_"):
        return f"Коммерческое предложение {key.rsplit('_', 1)[-1]}"
    return DOCUMENT_LABELS.get(key, document.display_name or key)


def _document_items(document) -> list:
    return [
        getattr(entry, "item", entry)
        for entry in (getattr(document, "items", []) or [])
    ]


def _normalized_item_name(item) -> str:
    return _value_text(item.name).lstrip("*").strip() or "Без наименования"


def _code_entries(document, field: str) -> list[str]:
    entries: list[str] = []
    seen: set[tuple[str, str]] = set()
    for item in _document_items(document):
        name = _normalized_item_name(item)
        for value in getattr(item, field, []) or []:
            code = _value_text(value)
            raw = str(value.raw_value or "")
            code_name = name
            if raw and raw != code:
                suffix = re.sub(
                    rf"^\s*{re.escape(code)}\s*[-–—]\s*",
                    "",
                    raw,
                ).strip()
                if suffix and suffix != raw:
                    code_name = suffix
            key = (code, code_name.casefold())
            if code and key not in seen:
                seen.add(key)
                entries.append(f"{code} - {code_name}")
    return entries


def _item_quantity_entries(document) -> list[str]:
    grouped: dict[str, list[str]] = defaultdict(list)
    labels: dict[str, str] = {}
    for item in _document_items(document):
        name = _normalized_item_name(item)
        key = name.casefold()
        labels.setdefault(key, name)
        quantity = _value_text(item.quantity)
        unit = _value_text(item.unit)
        if quantity:
            suffix = (
                f" {unit}"
                if unit and not quantity.casefold().endswith(unit.casefold())
                else ""
            )
            grouped[key].append(f"{quantity}{suffix}")
    return [
        f"{labels[key]} - {'; '.join(values)}"
        for key, values in grouped.items()
    ]


def _field_values(document, field: str) -> list[str]:
    value = getattr(document, field, None)
    values = value if isinstance(value, list) else [value]
    result = []
    seen = set()
    for entry in values:
        text = _value_text(entry)
        if not text:
            continue
        evidence = entry.evidence[0] if entry and entry.evidence else None
        key = (
            text.casefold(),
            evidence.table_id if evidence else None,
            evidence.row if evidence else None,
        )
        if key in seen:
            continue
        seen.add(key)
        result.append(text)
    return result


def _nmck_values(key: str, document) -> list[str]:
    values = _field_values(document, "price" if key == "contract" else "nmck")
    result = []
    for value in values:
        rendered = re.sub(r"\s*RUB\s*$", " рублей", value, flags=re.I)
        match = re.fullmatch(r"(\d+)(?:[.,]00)?(?:\s+рублей)?", rendered)
        if match:
            rendered = f"{int(match.group(1)):,}".replace(",", " ") + " рублей"
        elif "руб" not in rendered.casefold():
            rendered = f"{rendered} рублей"
        result.append(rendered)
    return result


def _append_occurrences(
    lines: list[str],
    title: str,
    package: ProcurementPackage,
    values: Callable[[str, object], list[str]],
    *,
    include: Iterable[str] = DOCUMENT_ORDER,
) -> None:
    included = set(include)
    lines.extend(["", title])
    any_document = False
    documents = dict(_documents(package))
    ordered_keys = list(include) + [
        key for key in documents if key not in included
    ]
    for key in ordered_keys:
        document = documents.get(key)
        if key not in included or document is None:
            continue
        any_document = True
        lines.append(f'Вхождения в документе "{_document_label(key, document)}":')
        entries = values(key, document)
        lines.extend(f"- {entry}" for entry in entries)
        if not entries:
            lines.append("- не найдено")
    if not any_document:
        lines.append("- документы отсутствуют")


def _display_with_labels(value, labels: dict[str, str]):
    if isinstance(value, dict):
        return {
            labels.get(str(key), str(key)): _display_with_labels(item, labels)
            for key, item in value.items()
        }
    if isinstance(value, list):
        return [_display_with_labels(item, labels) for item in value]
    return value


def _finding_details(
    finding: Finding,
    labels: dict[str, str],
    *,
    include_expected: bool = True,
) -> list[str]:
    marker = STATUS_MARKERS.get(finding.status, finding.status.value.upper())
    lines = [f"- {finding.title} — {marker}. {finding.message}".rstrip()]
    if include_expected and finding.expected is not None:
        lines.append(f"  Ожидалось: {_display_with_labels(finding.expected, labels)}")
    if finding.actual is not None:
        lines.append(f"  Получено: {_display_with_labels(finding.actual, labels)}")
    for evidence in finding.evidence[:5]:
        if not evidence.quote.strip():
            continue
        label = labels.get(evidence.document_id, evidence.document_id)
        location = []
        if evidence.table_id:
            location.append("таблица")
        if evidence.row is not None:
            location.append(f"строка {evidence.row + 1}")
        if evidence.column is not None:
            location.append(f"колонка {evidence.column + 1}")
        where = f", {', '.join(location)}" if location else ""
        lines.append(f"  Источник: {label}{where}: {evidence.quote}")
    return lines


def _append_findings(
    lines: list[str],
    findings: list[Finding],
    labels: dict[str, str],
    *,
    empty_message: str = "не обнаружены",
) -> None:
    relevant = [
        finding
        for finding in findings
        if finding.status != FindingStatus.PASSED
    ]
    lines.append("Ошибки и замечания:")
    if not relevant:
        lines.append(f"- {empty_message}")
        return
    for finding in relevant:
        lines.extend(_finding_details(finding, labels))


def _is_registry_characteristic(finding: Finding) -> bool:
    return finding.rule_id.startswith("registry.ktru.") and (
        ".characteristic." in finding.rule_id or ".required." in finding.rule_id
    )


def _is_registry_ktru(finding: Finding) -> bool:
    return finding.rule_id.startswith("registry.ktru.") and not _is_registry_characteristic(
        finding
    )


def _is_formatting_only_finding(finding: Finding) -> bool:
    text = f"{finding.title} {finding.message}".casefold()
    formatting_only = any(
        marker in text
        for marker in ("только в регистре", "регистр/оформлен", "регистр и оформлен")
    )
    no_conflict = any(
        marker in text
        for marker in (
            "противоречия не выявлено",
            "противоречий не выявлено",
            "смыслового расхождения не выявлено",
            "смысловых расхождений не выявлено",
            "нет смыслового противоречия",
        )
    )
    return formatting_only and no_conflict


def _is_technical_item_id_finding(finding: Finding) -> bool:
    if not finding.rule_id.startswith("items_consistency."):
        return False
    text = f"{finding.title} {finding.message}".casefold()
    return (
        "item_id" in text
        or "item-идентификатор" in text
        or ("техническ" in text and "идентификатор" in text)
    )


def _prepare_findings(findings: list[Finding]) -> list[Finding]:
    prepared = []
    for finding in findings:
        if _is_formatting_only_finding(finding) or _is_technical_item_id_finding(
            finding
        ):
            continue
        if (
            _is_registry_ktru(finding)
            and finding.status == FindingStatus.FAILED
            and "не удалось получить карточку" in finding.message.casefold()
        ):
            finding = finding.model_copy(
                update={
                    "status": FindingStatus.UNCERTAIN,
                    "severity": FindingSeverity.MANUAL_REVIEW,
                }
            )
        if (
            finding.source == "llm"
            and finding.status == FindingStatus.FAILED
            and "unresolved_fields"
            in f"{finding.title} {finding.message}".casefold()
        ):
            finding = finding.model_copy(
                update={
                    "status": FindingStatus.UNCERTAIN,
                    "severity": FindingSeverity.MANUAL_REVIEW,
                }
            )
        prepared.append(finding)
    return prepared


def _remaining_section(finding: Finding) -> str:
    rule = finding.rule_id.casefold()
    text = f"{finding.title} {finding.message}".casefold()
    if rule.startswith("delivery_and_finance."):
        return "Финансовые условия, НДС и порядок оплаты"
    if "security" in rule or "обеспеч" in text or "гарант" in text:
        return "Гарантии и обеспечения"
    if "smp" in rule or "сонко" in text or "смп" in text:
        return "Условия для СМП и СОНКО"
    if "national" in rule or "национальн" in text:
        return "Национальный режим"
    if "rights" in rule or "прав" in text or "обязанност" in text:
        return "Права и обязанности сторон"
    if "penalt" in rule or any(
        marker in text for marker in ("штраф", "пени", "неустой")
    ):
        return "Штрафы, пени и неустойки"
    return "Другие проверки"


def _supplier_price_lines(package: ProcurementPackage) -> list[str]:
    if package.onmck is None:
        return ["- документ ОНМЦК отсутствует"]
    lines = []
    for index, entry in enumerate(package.onmck.items, start=1):
        prices = [
            _value_text(price.unit_price)
            for price in entry.supplier_prices
            if _value_text(price.unit_price)
        ]
        coefficient = (
            f"{_clean_decimal(entry.variation_coefficient)}%"
            if entry.variation_coefficient is not None
            else "не рассчитан"
        )
        selected = _value_text(entry.selected_unit_price)
        details = f"цены поставщиков: {', '.join(prices) or 'не найдены'}"
        if selected:
            details += f"; выбранная цена: {selected}"
        details += f"; коэффициент вариации: {coefficient}"
        lines.append(
            f"{index}. {_normalized_item_name(entry.item)} — {details}"
        )
    return lines or ["- позиции не извлечены"]


def build_report_text(
    findings: list[Finding],
    *,
    package: ProcurementPackage | None = None,
    detailed: bool = True,
    document_labels: dict[str, str] | None = None,
) -> str:
    findings = _prepare_findings(findings)
    labels = document_labels or {}
    counts = Counter(finding.status for finding in findings)
    lines = [
        "Результат проверки закупочной документации",
        (
            f"Ошибок: {counts[FindingStatus.FAILED]}. "
            f"Требуют проверки: {counts[FindingStatus.UNCERTAIN]}. "
            f"Успешных проверок: {counts[FindingStatus.PASSED]}. "
            f"Пропущено: {counts[FindingStatus.SKIPPED]}."
        ),
    ]

    if package is None:
        grouped: dict[str, list[Finding]] = defaultdict(list)
        for finding in findings:
            title = "Прочие проверки"
            for prefix, candidate in FALLBACK_SECTIONS:
                if finding.rule_id.startswith(prefix):
                    title = candidate
                    break
            grouped[title].append(finding)
        section_order = [title for _, title in FALLBACK_SECTIONS] + ["Прочие проверки"]
        section_number = 0
        for title in dict.fromkeys(section_order):
            entries = grouped.get(title)
            if not entries:
                continue
            lines.extend(["", f"{section_number}) {title}"])
            section_number += 1
            for finding in entries:
                if finding.status == FindingStatus.PASSED:
                    lines.append(f"- {finding.title} — ОК")
                else:
                    lines.extend(_finding_details(finding, labels))
        return "\n".join(lines)

    loaded = [_document_label(key, document) for key, document in _documents(package)]
    missing = [
        DOCUMENT_LABELS[key]
        for key in DOCUMENT_ORDER
        if getattr(package, key, None) is None
    ]
    lines.extend(
        [
            "",
            "0) Комплектность пакета",
            "Загруженные документы:",
            *(f"- {label}" for label in loaded),
            "Отсутствующие документы:",
            *(f"- {label}" for label in missing),
        ]
    )
    if not missing:
        lines.append("- нет")
    _append_findings(
        lines,
        [finding for finding in findings if finding.rule_id.startswith("package.")],
        labels,
    )

    ktru_findings = [finding for finding in findings if _is_registry_ktru(finding)]
    lines.extend(["", "1) Проверка КТРУ через сервис zakupki.gov.ru"])
    for finding in ktru_findings:
        if finding.status == FindingStatus.PASSED:
            lines.append(f"- {finding.title} — карточка получена")
        else:
            lines.extend(_finding_details(finding, labels, include_expected=False))
    if not ktru_findings:
        lines.append("- коды КТРУ не найдены")

    okpd_findings = [
        finding
        for finding in findings
        if finding.rule_id.startswith("registry.okpd2.")
    ]
    lines.extend(["", "2) Проверка ОКПД2 по ПП №1875"])
    for finding in okpd_findings:
        lines.append(f"- {finding.title}: {finding.message}")
    if not okpd_findings:
        lines.append("- коды ОКПД2 не найдены")

    lines.extend(["", "3) Внутренний анализ перечня документов"])
    _append_occurrences(
        lines,
        "Проверка ОКПД2:",
        package,
        lambda _key, document: _code_entries(document, "okpd2"),
        include=("plan", "contract", "ooz"),
    )
    _append_findings(
        lines,
        [finding for finding in findings if finding.rule_id == "items.okpd2_consistency"],
        labels,
    )
    _append_occurrences(
        lines,
        "Проверка КТРУ:",
        package,
        lambda _key, document: _code_entries(document, "ktru"),
        include=("plan", "contract", "ooz"),
    )
    _append_findings(
        lines,
        [finding for finding in findings if finding.rule_id == "items.ktru_consistency"],
        labels,
    )
    _append_occurrences(
        lines,
        "Проверка наименований и количества товаров:",
        package,
        lambda _key, document: _item_quantity_entries(document),
        include=("plan", "explanatory_note", "ooz", "contract", "onmck"),
    )
    _append_findings(
        lines,
        [
            finding
            for finding in findings
            if finding.rule_id.startswith("items.")
            and finding.rule_id
            not in {"items.okpd2_consistency", "items.ktru_consistency"}
        ],
        labels,
    )
    _append_occurrences(
        lines,
        "Проверка начальной (максимальной) цены контракта:",
        package,
        _nmck_values,
        include=("plan", "request", "onmck", "explanatory_note", "contract"),
    )
    _append_occurrences(
        lines,
        "Проверка сроков поставки:",
        package,
        lambda _key, document: _field_values(document, "delivery_periods"),
        include=("plan", "request", "explanatory_note", "ooz", "contract"),
    )
    _append_occurrences(
        lines,
        "Проверка места поставки:",
        package,
        lambda _key, document: _field_values(document, "delivery_places"),
        include=("plan", "request", "explanatory_note", "ooz", "contract"),
    )

    characteristic_findings = [
        finding for finding in findings if _is_registry_characteristic(finding)
    ]
    lines.extend(["", "4) Внешняя проверка характеристик ООЗ по КТРУ"])
    if characteristic_findings:
        passed = [
            finding
            for finding in characteristic_findings
            if finding.status == FindingStatus.PASSED
        ]
        if passed:
            lines.append(f"- успешно проверено характеристик: {len(passed)}")
        for finding in characteristic_findings:
            if finding.status != FindingStatus.PASSED:
                lines.extend(_finding_details(finding, labels))
    elif any(finding.status != FindingStatus.PASSED for finding in ktru_findings):
        lines.append(
            "- проверка не выполнена: карточки КТРУ не были получены с внешнего сайта"
        )
    else:
        lines.append("- характеристики для проверки не найдены")

    lines.extend(["", "5) Сравнение цен поставщиков в ОНМЦК"])
    lines.extend(_supplier_price_lines(package))
    _append_findings(
        lines,
        [finding for finding in findings if finding.rule_id.startswith("price.")],
        labels,
    )

    consumed = {
        id(finding)
        for finding in findings
        if (
            finding.rule_id.startswith("package.")
            or finding.rule_id.startswith("registry.")
            or finding.rule_id.startswith("items.")
            or finding.rule_id.startswith("price.")
        )
    }
    remaining = [
        finding
        for finding in findings
        if id(finding) not in consumed and finding.status != FindingStatus.PASSED
    ]
    lines.extend(["", "6) Дополнительные проверки"])
    grouped_remaining: dict[str, list[Finding]] = defaultdict(list)
    for finding in remaining:
        grouped_remaining[_remaining_section(finding)].append(finding)
    section_order = (
        "Финансовые условия, НДС и порядок оплаты",
        "Гарантии и обеспечения",
        "Условия для СМП и СОНКО",
        "Национальный режим",
        "Права и обязанности сторон",
        "Штрафы, пени и неустойки",
        "Другие проверки",
    )
    if not remaining:
        lines.append("- замечаний не обнаружено")
    for title in section_order:
        entries = grouped_remaining.get(title)
        if not entries:
            continue
        lines.extend(["", f"{title}:"])
        _append_findings(lines, entries, labels)
    return "\n".join(lines)


def build_report_docx_bytes(
    findings: list[Finding],
    *,
    package: ProcurementPackage | None = None,
    detailed: bool = True,
    document_labels: dict[str, str] | None = None,
) -> bytes:
    text = build_report_text(
        findings,
        package=package,
        detailed=detailed,
        document_labels=document_labels,
    )
    document = Document()
    for index, line in enumerate(text.splitlines()):
        stripped = line.strip()
        if not stripped:
            continue
        if index == 0:
            document.add_heading(stripped, level=1)
            continue
        if stripped[0:1].isdigit() and ") " in stripped[:4]:
            document.add_heading(stripped, level=2)
            continue
        if stripped.endswith(":") and not stripped.startswith("-"):
            document.add_heading(stripped[:-1], level=3)
            continue
        if stripped.startswith("- "):
            paragraph = document.add_paragraph(style="List Bullet")
            content = stripped[2:]
        else:
            paragraph = document.add_paragraph()
            content = stripped
            if line.startswith("  "):
                paragraph.paragraph_format.left_indent = Inches(0.35)
        paragraph.paragraph_format.space_after = Pt(3)
        run = paragraph.add_run(content)
        if "— ОШИБКА" in content:
            run.font.color.rgb = STATUS_COLORS[FindingStatus.FAILED]
        elif "ТРЕБУЕТ ПРОВЕРКИ" in content:
            run.font.color.rgb = STATUS_COLORS[FindingStatus.UNCERTAIN]
        elif "ПРОВЕРКА ПРОПУЩЕНА" in content:
            run.font.color.rgb = STATUS_COLORS[FindingStatus.SKIPPED]

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()
