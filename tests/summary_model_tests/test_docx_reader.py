from docx import Document

from summary_model.extraction.llm_client import render_ir_chunks
from summary_model.ingestion import read_docx


def test_reader_preserves_block_order_and_merged_cells(tmp_path):
    path = tmp_path / "sample.docx"
    document = Document()
    document.add_paragraph("Before table")
    table = document.add_table(rows=2, cols=3)
    table.cell(0, 0).text = "Header"
    table.cell(0, 0).merge(table.cell(0, 1))
    table.cell(0, 2).text = "Value"
    table.cell(1, 0).text = "1"
    table.cell(1, 1).text = "Item"
    table.cell(1, 2).text = "10"
    document.add_paragraph("After table")
    document.save(path)

    result = read_docx(path)

    nonempty = [
        block for block in result.blocks
        if block.type == "table" or block.text
    ]
    assert [block.type for block in nonempty] == ["paragraph", "table", "paragraph"]
    assert nonempty[0].text == "Before table"
    assert nonempty[2].text == "After table"

    table_ir = nonempty[1].table
    assert table_ir is not None
    assert table_ir.row_count == 2
    assert table_ir.column_count == 3
    assert table_ir.origin_value(0, 0) == "Header"
    assert table_ir.span_at_origin(0, 0) == (1, 2)
    assert table_ir.matrix()[0][:2] == ["Header", "Header"]
    first_row = table_ir.rows[0]
    assert first_row.values == {"c0": "Header", "c2": "Value"}
    assert first_row.spans == {"c0": (1, 2)}
    assert str(table_ir.rows).count("Header") == 1
    rendered = "\n".join(render_ir_chunks(result))
    assert rendered.count("Header") == 1
    assert '"row"' not in rendered
    assert '"column"' not in rendered
    assert '"text"' not in rendered


def test_reader_keeps_vertical_merge_as_one_origin_cell(tmp_path):
    path = tmp_path / "vertical.docx"
    document = Document()
    table = document.add_table(rows=3, cols=2)
    table.cell(0, 0).text = "Item"
    table.cell(0, 0).merge(table.cell(2, 0))
    table.cell(0, 1).text = "A"
    table.cell(1, 1).text = "B"
    table.cell(2, 1).text = "C"
    document.save(path)

    result = read_docx(path)
    table_ir = next(block.table for block in result.blocks if block.table)
    assert table_ir.origin_value(0, 0) == "Item"
    assert table_ir.span_at_origin(0, 0) == (3, 1)
    assert [row[0] for row in table_ir.matrix()] == ["Item", "Item", "Item"]


def test_reader_attaches_table_context(tmp_path):
    path = tmp_path / "context.docx"
    document = Document()
    document.add_paragraph("Таблица №1 Характеристики товара")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Характеристика"
    table.cell(0, 1).text = "Значение характеристики"
    table.cell(1, 0).text = "Цвет"
    table.cell(1, 1).text = "Черный"
    document.save(path)

    result = read_docx(path)
    table_ir = next(block.table for block in result.blocks if block.table)
    assert table_ir.title == "Таблица №1 Характеристики товара"
    assert table_ir.kind == "characteristics"


def test_reader_omits_empty_paragraph_blocks(tmp_path):
    path = tmp_path / "empty-paragraphs.docx"
    document = Document()
    document.add_paragraph("")
    document.add_paragraph("   ")
    document.add_paragraph("Meaningful")
    document.add_paragraph("")
    document.save(path)

    result = read_docx(path)

    paragraphs = [block for block in result.blocks if block.type == "paragraph"]
    assert [block.text for block in paragraphs] == ["Meaningful"]
