from docx import Document

from summary_model.classification import DocumentClassifier
from summary_model.domain.models import DocumentType, PlanRequestSummary
from summary_model.extraction.llm_client import StructuredLLMClient
from summary_model.ingestion import read_docx


class FakeRunnable:
    def __init__(self):
        self.calls = 0

    def invoke(self, _prompt):
        self.calls += 1
        if self.calls == 1:
            raise ValueError("invalid first response")
        return {
            "detected_type": "plan",
            "classification_confidence": 0.9,
            "subject": {
                "raw_value": "Поставка шин",
                "normalized_value": "Поставка шин",
                "confidence": 0.9,
                "evidence": [],
                "warnings": [],
            },
        }


class FakeModel:
    def __init__(self):
        self.runnable = FakeRunnable()
        self.structured_output_method = None

    def with_structured_output(self, _schema, *, method):
        self.structured_output_method = method
        return self.runnable


def test_classifier_uses_content_and_reports_hint_conflict(tmp_path):
    path = tmp_path / "document.docx"
    document = Document()
    document.add_paragraph("ОПИСАНИЕ ОБЪЕКТА ЗАКУПКИ")
    document.add_paragraph("Функциональные, технические и качественные характеристики")
    document.save(path)
    ir = read_docx(path)

    result = DocumentClassifier().classify(ir, DocumentType.CONTRACT)

    assert result.document_type == DocumentType.OOZ
    assert result.warnings


def test_structured_client_retries_once():
    model = FakeModel()
    client = StructuredLLMClient(model=model)

    result, error = client.extract(PlanRequestSummary, "Extract plan", "payload")

    assert error is None
    assert result is not None
    assert result.subject.normalized_value == "Поставка шин"
    assert model.runnable.calls == 2
    assert model.structured_output_method == "function_calling"
