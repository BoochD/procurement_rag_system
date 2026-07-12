import json

from docx import Document

from summary_model.domain.models import DocumentType, InputDocument
from summary_model.extraction_cli import main as extraction_cli_main
from summary_model.extraction_pipeline import extract_package
from summary_model.ingestion import read_docx
from summary_model.tables import extract_tables


def _save_plan(path):
    document = Document()
    document.add_paragraph("Заявка на внесение в план-график")
    table = document.add_table(rows=5, cols=2)
    rows = [
        ("Наименование объекта закупки", "Поставка картриджей"),
        ("Код позиции КТРУ", "20.59.12.120-00000002 - Картридж"),
        ("Начальная максимальная цена контракта", "350 000 рублей 00 копеек"),
        ("Преимущества СМП", "нет"),
        ("Количество", "10 шт"),
    ]
    for index, (key, value) in enumerate(rows):
        table.cell(index, 0).text = key
        table.cell(index, 1).text = value
    signatures = document.add_table(rows=3, cols=2)
    signatures.cell(0, 0).text = "Иванов И.И."
    signatures.cell(1, 0).text = "Петров П.П."
    signatures.cell(2, 0).text = "Сидоров С.С."
    document.save(path)


def _save_plan_with_duplicate_value_columns(path):
    document = Document()
    document.add_paragraph("Заявка на внесение в план-график")
    table = document.add_table(rows=2, cols=4)
    rows = [
        ("", "Наименование объекта закупки", "Поставка шин", "Поставка шин"),
        ("", "Количество", "39 шт", "39 шт"),
    ]
    for row_index, values in enumerate(rows):
        for col_index, value in enumerate(values):
            table.cell(row_index, col_index).text = value
    document.save(path)


def _save_plan_with_distinct_value_columns(path):
    document = Document()
    document.add_paragraph("Заявка на внесение в план-график")
    table = document.add_table(rows=1, cols=4)
    values = ("", "Применение национального режима", "Нет", "Основание неприменения")
    for col_index, value in enumerate(values):
        table.cell(0, col_index).text = value
    document.save(path)


def _save_request_with_text_attachment_table(path):
    document = Document()
    document.add_paragraph("Обращение о проведении закупки")
    table = document.add_table(rows=5, cols=2)
    rows = [
        ("Приложение:", "1. Заявка на внесение в план-график;"),
        ("", "2. Определение цены контракта с единственным исполнителем;"),
        ("", "3. Проект контракта;"),
        ("", "4. Описание объекта закупки;"),
        ("", "5. Пояснительная записка."),
    ]
    for row_index, values in enumerate(rows):
        for col_index, value in enumerate(values):
            table.cell(row_index, col_index).text = value
    document.save(path)


def _save_request_with_plain_text_attachment_lines(path):
    document = Document()
    document.add_paragraph("Обращение о проведении закупки")
    document.add_paragraph("Приложение:")
    document.add_paragraph("Заявка на внесение в план-график")
    document.add_paragraph("Определение цены контракта с единственным исполнителем")
    document.add_paragraph("Проект контракта")
    document.add_paragraph("Описание объекта закупки")
    document.add_paragraph("Пояснительная записка")
    document.save(path)


def _save_ooz(path):
    document = Document()
    document.add_paragraph("Описание объекта закупки")
    table = document.add_table(rows=3, cols=7)
    headers = [
        "№",
        "Наименование товара",
        "ОКПД2 / КТРУ",
        "Наименование характеристики",
        "Значение характеристики",
        "Единица измерения товара",
        "Количество",
    ]
    for index, header in enumerate(headers):
        table.cell(0, index).text = header
    values = [
        "1",
        "Картридж",
        "20.59.12.120-00000002",
        "Цвет",
        "Черный",
        "шт",
        "10",
    ]
    for index, value in enumerate(values):
        table.cell(1, index).text = value
    table.cell(2, 3).text = "Ресурс, страниц"
    table.cell(2, 4).text = ">= 8000"
    document.save(path)


def _save_combined_ooz(path):
    document = Document()
    document.add_paragraph("Проект контракта")
    document.add_paragraph("ОПИСАНИЕ ОБЪЕКТА ЗАКУПКИ")
    table = document.add_table(rows=4, cols=5)
    headers = [
        "Наименование, ОКПД2/КТРУ",
        "Наименование характеристики",
        "Значение характеристики",
        "Единица измерения характеристики",
        "Количество, штук",
    ]
    rows = [
        headers,
        [
            "Программное обеспечение, КТРУ 58.29.11.000-00000003",
            "Класс программ для электронных вычислительных машин и баз данных",
            "(12.10) Программное обеспечение для решения отраслевых задач",
            "",
            "506",
        ],
        [
            "Программное обеспечение, КТРУ 58.29.11.000-00000003",
            "Способ предоставления",
            "Копия электронного экземпляра",
            "",
            "506",
        ],
        [
            "Программное обеспечение, КТРУ 58.29.11.000-00000003",
            "Вид лицензии",
            "Простая (неисключительная)",
            "",
            "506",
        ],
    ]
    for row_index, values in enumerate(rows):
        for col_index, value in enumerate(values):
            table.cell(row_index, col_index).text = value
    document.save(path)


def _save_onmck(path):
    document = Document()
    document.add_paragraph("Обоснование начальной максимальной цены контракта")
    table = document.add_table(rows=2, cols=12)
    headers = [
        "№",
        "Наименование товара",
        "Единица измерения",
        "Количество",
        "Поставщик 1 цена за ед. товара",
        "Поставщик 1 стоимость товаров",
        "Поставщик 2 цена за ед. товара",
        "Поставщик 2 стоимость товаров",
        "Поставщик 3 цена за ед. товара",
        "Поставщик 3 стоимость товаров",
        "Минимальная цена",
        "Цена контракта",
    ]
    for index, header in enumerate(headers):
        table.cell(0, index).text = header
    values = ["1", "Картридж", "шт", "10", "100", "1000", "90", "900", "110", "1100", "90", "900"]
    for index, value in enumerate(values):
        table.cell(1, index).text = value
    document.save(path)


def _save_onmck_with_executors(path):
    document = Document()
    document.add_paragraph("Обоснование начальной максимальной цены контракта")
    table = document.add_table(rows=5, cols=12)
    rows = [
        [
            "№ п/п",
            "Наименование",
            "Ед. изм.",
            "Кол-во",
            "Цена услуги (руб.)/источники информации о ценах",
            "Цена услуги (руб.)/источники информации о ценах",
            "Цена услуги (руб.)/источники информации о ценах",
            "Цена услуги (руб.)/источники информации о ценах",
            "Цена услуги (руб.)/источники информации о ценах",
            "Цена услуги (руб.)/источники информации о ценах",
            "Минимальная цена за ед. (руб.)",
            "Начальная (максимальная) цена контракта (руб.)",
        ],
        [
            "№ п/п",
            "Наименование",
            "Ед. изм.",
            "Кол-во",
            "Исполнитель 1",
            "Исполнитель 1",
            "Исполнитель 2",
            "Исполнитель 2",
            "Исполнитель 3",
            "Исполнитель 3",
            "Минимальная цена за ед. (руб.)",
            "Начальная (максимальная) цена контракта (руб.)",
        ],
        [
            "№ п/п",
            "Наименование",
            "Ед. изм.",
            "Кол-во",
            "Цена за ед., руб.",
            "Стоимость, руб.",
            "Цена за ед., руб.",
            "Стоимость, руб.",
            "Цена за ед., руб.",
            "Стоимость, руб.",
            "Минимальная цена за ед. (руб.)",
            "Начальная (максимальная) цена контракта (руб.)",
        ],
        [
            "1",
            "Программное обеспечение, КТРУ 58.29.11.000-00000003",
            "Шт.",
            "506",
            "4 461,00",
            "2 257 266,00",
            "4 377,00",
            "2 214 762,00",
            "4 169,00",
            "2 109 514,00",
            "4 169,00",
            "2 109 514,00",
        ],
        [
            "",
            "ИТОГО",
            "",
            "",
            "2 257 266,00",
            "2 257 266,00",
            "2 214 762,00",
            "2 214 762,00",
            "2 109 514,00",
            "2 109 514,00",
            "",
            "2 109 514,00",
        ],
    ]
    for row_index, values in enumerate(rows):
        for col_index, value in enumerate(values):
            table.cell(row_index, col_index).text = value
    document.save(path)


def _save_contract(path):
    document = Document()
    document.add_paragraph("Проект контракта")
    ooz_table = document.add_table(rows=2, cols=7)
    ooz_headers = [
        "№",
        "Наименование товара",
        "ОКПД2 / КТРУ",
        "Наименование характеристики",
        "Значение характеристики",
        "Единица измерения товара",
        "Количество",
    ]
    for index, header in enumerate(ooz_headers):
        ooz_table.cell(0, index).text = header
    ooz_values = [
        "1",
        "Картридж",
        "20.59.12.120-00000002",
        "Цвет",
        "Черный",
        "шт",
        "10",
    ]
    for index, value in enumerate(ooz_values):
        ooz_table.cell(1, index).text = value

    document.add_paragraph("Спецификация")
    spec = document.add_table(rows=3, cols=10)
    spec_headers = [
        "Наименование продукции",
        "Характеристики продукции",
        "Ед. изм.",
        "Кол-во",
        "Цена за единицу без НДС",
        "Сумма без НДС",
        "Ставка НДС",
        "Сумма НДС",
        "Цена за единицу с учетом НДС",
        "Всего, рублей",
    ]
    for index, header in enumerate(spec_headers):
        spec.cell(0, index).text = header
    spec_values = [
        "Картридж",
        "Черный",
        "шт",
        "10",
        "90",
        "900",
        "-",
        "-",
        "90",
        "900",
    ]
    for index, value in enumerate(spec_values):
        spec.cell(1, index).text = value
    for index in range(10):
        spec.cell(2, index).text = "900"
    spec.cell(2, 0).text = "Всего:"
    document.save(path)


def _save_template_specification_contract(path):
    document = Document()
    document.add_paragraph("Проект контракта")
    document.add_paragraph("Спецификация")
    empty_spec = document.add_table(rows=2, cols=8)
    headers = [
        "№ п/п",
        "Наименование программного обеспечения",
        "Единица измерения",
        "Количество лицензий",
        "Действует с",
        "Действует по",
        "Цена единицы без НДС",
        "Сумма без НДС",
    ]
    for index, header in enumerate(headers):
        empty_spec.cell(0, index).text = header
    for index in range(8):
        empty_spec.cell(1, index).text = ""

    numbered_spec = document.add_table(rows=3, cols=8)
    for index, header in enumerate(headers):
        numbered_spec.cell(0, index).text = header
    for index in range(8):
        numbered_spec.cell(1, index).text = str(index + 1)
    numbered_spec.cell(2, 0).text = "Всего:"
    numbered_spec.cell(2, 7).text = "0"
    document.save(path)


def _save_contract_with_text_attachments_only(path):
    document = Document()
    document.add_paragraph("Проект контракта")
    document.add_paragraph("8. Обеспечение исполнения Контракта")
    document.add_paragraph("8.1. Обеспечение исполнения контракта не предусмотрено.")
    document.add_paragraph("9. Срок исполнения Контракта")
    document.add_paragraph("9.2. Срок исполнения Контракта: 70 календарных дней.")
    document.add_paragraph("14. Приложения")
    document.add_paragraph("14.1. Неотъемлемыми частями Контракта являются следующие приложения:")
    document.add_paragraph("приложение № 1 «Описание объекта закупки»;")
    document.add_paragraph("приложение № 2 «Акт приема-передачи товара» (форма);")
    document.add_paragraph("приложение № 3 «Спецификация».")
    document.save(path)


def test_key_value_table_preserves_all_raw_fields_and_negative_values(tmp_path):
    path = tmp_path / "plan.docx"
    _save_plan(path)
    package = extract_package(
        [InputDocument(path=path, type_hint=DocumentType.PLAN, display_name="plan")]
    )

    plan = package.schedule_application
    assert plan is not None
    assert plan.purchase_subject == "Поставка картриджей"
    assert plan.ktru_codes == ["20.59.12.120-00000002"]
    assert plan.nmck and plan.nmck.amount == 350000
    assert len(plan.raw_fields) == 5
    assert "Преимущества СМП" in plan.negative_value_fields
    assert all("Иванов" not in field.key for field in plan.raw_fields)

    ir = read_docx(path)
    tables = extract_tables(ir, DocumentType.PLAN)
    assert [table.table_type for table in tables] == [
        "schedule_application_table",
        "signature_table",
    ]


def test_key_value_table_deduplicates_identical_value_columns(tmp_path):
    path = tmp_path / "plan_duplicate_columns.docx"
    _save_plan_with_duplicate_value_columns(path)
    package = extract_package(
        [InputDocument(path=path, type_hint=DocumentType.PLAN, display_name="plan")]
    )

    plan = package.schedule_application
    assert plan is not None
    assert plan.purchase_subject == "Поставка шин"
    assert plan.raw_fields_dict["Количество"] == "39 шт"
    assert "; 39 шт" not in plan.raw_fields_dict["Количество"]


def test_key_value_table_preserves_distinct_value_columns(tmp_path):
    path = tmp_path / "plan_distinct_columns.docx"
    _save_plan_with_distinct_value_columns(path)
    package = extract_package(
        [InputDocument(path=path, type_hint=DocumentType.PLAN, display_name="plan")]
    )

    plan = package.schedule_application
    assert plan is not None
    assert plan.raw_fields_dict["Применение национального режима"] == "Нет; Основание неприменения"


def test_purchase_request_extracts_text_attachment_list_from_small_table(tmp_path):
    path = tmp_path / "request_attachments.docx"
    _save_request_with_text_attachment_table(path)
    package = extract_package(
        [InputDocument(path=path, type_hint=DocumentType.REQUEST, display_name="request")]
    )

    request = package.purchase_request
    assert request is not None
    assert [
        (item.number, item.normalized_document_type)
        for item in request.attachments
    ] == [
        ("1", "schedule_application"),
        ("2", "nmck_justification"),
        ("3", "contract_draft"),
        ("4", "purchase_description"),
        ("5", "explanatory_note"),
    ]


def test_purchase_request_extracts_plain_text_attachment_lines_without_semicolons(tmp_path):
    path = tmp_path / "request_plain_attachments.docx"
    _save_request_with_plain_text_attachment_lines(path)
    package = extract_package(
        [InputDocument(path=path, type_hint=DocumentType.REQUEST, display_name="request")]
    )

    request = package.purchase_request
    assert request is not None
    assert [
        (item.number, item.normalized_document_type)
        for item in request.attachments
    ] == [
        ("1", "schedule_application"),
        ("2", "nmck_justification"),
        ("3", "contract_draft"),
        ("4", "purchase_description"),
        ("5", "explanatory_note"),
    ]


def test_ooz_logical_rows_attach_characteristics_to_items(tmp_path):
    path = tmp_path / "ooz.docx"
    _save_ooz(path)
    ir = read_docx(path)
    tables = extract_tables(ir, DocumentType.OOZ)
    table = tables[0]

    assert table.table_type == "ooz_items_table"
    assert [row.row_type for row in table.logical_rows] == [
        "item",
        "characteristic",
        "characteristic",
    ]
    item = table.compact_json["items"][0]
    assert item["name"] == "Картридж"
    assert item["ktru_code"] == "20.59.12.120-00000002"
    assert item["unit"] == "шт"
    assert item["quantity"] == 10
    assert [char["name"] for char in item["characteristics"]] == [
        "Цвет",
        "Ресурс, страниц",
    ]


def test_ooz_combined_name_and_ktru_rows_group_into_one_item(tmp_path):
    path = tmp_path / "combined_ooz.docx"
    _save_combined_ooz(path)
    package = extract_package(
        [InputDocument(path=path, type_hint=DocumentType.CONTRACT, display_name="contract")]
    )

    contract = package.contract_draft
    assert contract is not None
    assert len(contract.items) == 1
    item = contract.items[0]
    assert item.name == "Программное обеспечение"
    assert item.ktru_code == "58.29.11.000-00000003"
    assert item.quantity == 506
    assert len(item.characteristics) == 3
    assert [characteristic.name for characteristic in item.characteristics] == [
        "Класс программ для электронных вычислительных машин и баз данных",
        "Способ предоставления",
        "Вид лицензии",
    ]


def test_onmck_table_extracts_supplier_prices_and_recalculates(tmp_path):
    path = tmp_path / "onmck.docx"
    _save_onmck(path)
    package = extract_package(
        [InputDocument(path=path, type_hint=DocumentType.ONMCK, display_name="onmck")]
    )

    onmck = package.nmck_justification
    assert onmck is not None
    assert len(onmck.items) == 1
    item = onmck.items[0]
    assert len(onmck.price_sources) == 3
    assert [price.unit_price for price in item.supplier_prices] == [100, 90, 110]
    assert [price.row_total for price in item.supplier_prices] == [1000, 900, 1100]
    assert item.quantity == 10
    assert item.selected_min_unit_price == 90
    assert item.row_total_declared == 900
    assert item.calculated_min_unit_price == 90
    assert item.row_total_calculated == 900
    assert item.is_declared_min_price_correct is True
    assert item.is_row_total_correct is True


def test_onmck_table_accepts_executor_price_sources(tmp_path):
    path = tmp_path / "onmck_executors.docx"
    _save_onmck_with_executors(path)
    package = extract_package(
        [InputDocument(path=path, type_hint=DocumentType.ONMCK, display_name="onmck")]
    )

    onmck = package.nmck_justification
    assert onmck is not None
    assert len(onmck.items) == 1
    assert len(onmck.price_sources) == 3
    item = onmck.items[0]
    assert item.name == "Программное обеспечение, КТРУ 58.29.11.000-00000003"
    assert item.quantity == 506
    assert [price.unit_price for price in item.supplier_prices] == [4461, 4377, 4169]
    assert [price.row_total for price in item.supplier_prices] == [
        2257266,
        2214762,
        2109514,
    ]
    assert item.selected_min_unit_price == 4169
    assert item.row_total_declared == 2109514
    assert item.is_declared_min_price_correct is True
    assert item.is_row_total_correct is True


def test_contract_keeps_description_and_specification_separate(tmp_path):
    path = tmp_path / "contract.docx"
    _save_contract(path)
    package = extract_package(
        [InputDocument(path=path, type_hint=DocumentType.CONTRACT, display_name="contract")]
    )

    contract = package.contract_draft
    assert contract is not None
    assert len(contract.items) == 1
    assert contract.items[0].name == "Картридж"
    assert len(contract.specification_items) == 1
    specification = contract.specification_items[0]
    assert specification.name == "Картридж"
    assert specification.quantity == 10
    assert specification.total_price == 900
    assert all(item.name != "Всего:" for item in contract.specification_items)


def test_contract_template_specification_does_not_create_fake_items(tmp_path):
    path = tmp_path / "template_contract.docx"
    _save_template_specification_contract(path)
    package = extract_package(
        [InputDocument(path=path, type_hint=DocumentType.CONTRACT, display_name="contract")]
    )

    contract = package.contract_draft
    assert contract is not None
    assert contract.specification_items == []
    tables = extract_tables(read_docx(path), DocumentType.CONTRACT)
    spec_tables = [table for table in tables if table.table_type == "contract_specification_table"]
    assert spec_tables
    assert all(not table.compact_json.get("items") for table in spec_tables)
    assert any(table.compact_json.get("totals") for table in spec_tables)


def test_contract_extracts_text_referenced_attachments_and_warns_when_tables_missing(tmp_path):
    path = tmp_path / "contract_text_attachments.docx"
    _save_contract_with_text_attachments_only(path)
    package = extract_package(
        [InputDocument(path=path, type_hint=DocumentType.CONTRACT, display_name="contract")]
    )

    contract = package.contract_draft
    assert contract is not None
    assert [
        (item.number, item.title_raw, item.normalized_document_type, item.attachment_kind)
        for item in contract.referenced_attachments
    ] == [
        ("1", "Описание объекта закупки", "purchase_description", "purchase_description"),
        ("2", "Акт приема-передачи товара", "unknown", "acceptance_act_form"),
        ("3", "Спецификация", "unknown", "contract_specification"),
    ]
    assert contract.actual_attachments == []
    assert contract.contract_security is not None
    assert contract.contract_security.is_not_required is True
    assert contract.contract_execution_term_text == "70 календарных дней."
    assert contract.contract_execution_term is not None
    assert contract.contract_execution_term.days == 70
    assert contract.items == []
    assert contract.specification_items == []
    assert any("Описание объекта закупки" in warning for warning in contract.parser_warnings)
    assert any("Спецификация" in warning for warning in contract.parser_warnings)


def test_extraction_cli_creates_result_tables_and_debug_artifacts(tmp_path):
    input_dir = tmp_path / "input"
    output_dir = tmp_path / "output"
    input_dir.mkdir()
    _save_plan(input_dir / "plan.docx")

    exit_code = extraction_cli_main(
        [
            "--input-dir",
            str(input_dir),
            "--output-dir",
            str(output_dir),
        ]
    )

    assert exit_code == 0
    assert (output_dir / "extraction_result.json").exists()
    assert list((output_dir / "documents").glob("*.json"))
    assert list((output_dir / "tables").glob("*.json"))
    assert list((output_dir / "llm_payloads").glob("*.json"))
    assert list((output_dir / "debug" / "tables").glob("*/*_physical.md"))
    assert list((output_dir / "debug" / "tables").glob("*/*_logical.json"))
    assert list((output_dir / "debug" / "tables").glob("*/*_compact.md"))
    payload = json.loads((output_dir / "extraction_result.json").read_text(encoding="utf-8"))
    assert payload["schema_version"] == "extraction-1.1.0"
    assert payload["schedule_application"]["purchase_subject"] == "Поставка картриджей"
    compact = next((output_dir / "debug" / "tables").glob("*/*_compact.md")).read_text(
        encoding="utf-8"
    )
    assert '"row"' not in compact
    assert "TABLE schedule_application_table" in compact
    llm_payload = json.loads(
        next((output_dir / "llm_payloads").glob("*.json")).read_text(encoding="utf-8")
    )
    assert llm_payload["schema_version"] == "llm-payload-1.0.0"
    assert llm_payload["plain_text_blocks"]
    assert llm_payload["known_extracted"]["document_type"] == "schedule_application"
    assert [table["table_type"] for table in llm_payload["tables"]] == [
        "schedule_application_table"
    ]
    assert "compact_json" not in llm_payload["tables"][0]
    assert "compact_markdown" not in llm_payload["tables"][0]
