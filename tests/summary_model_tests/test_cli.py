import json

from docx import Document

from summary_model.cli import main


def _write_plan(path):
    document = Document()
    document.add_paragraph("Заявка на внесение в план-график")
    table = document.add_table(rows=3, cols=2)
    table.cell(0, 0).text = "Наименование объекта закупки"
    table.cell(0, 1).text = "Поставка шин"
    table.cell(1, 0).text = "Код позиции КТРУ"
    table.cell(1, 1).text = "22.11.11.000-00000007 - Шина"
    table.cell(2, 0).text = "Начальная максимальная цена контракта"
    table.cell(2, 1).text = "350 000 рублей"
    document.save(path)


def test_cli_creates_expected_artifacts(tmp_path):
    input_dir = tmp_path / "input"
    output_dir = tmp_path / "output"
    input_dir.mkdir()
    _write_plan(input_dir / "plan.docx")

    exit_code = main(
        [
            "--input-dir",
            str(input_dir),
            "--output-dir",
            str(output_dir),
            "--no-llm",
            "--no-external",
        ]
    )

    assert exit_code == 0
    for name in ("package.json", "findings.json", "report.txt", "report.docx", "run.json"):
        assert (output_dir / name).exists()
    run = json.loads((output_dir / "run.json").read_text(encoding="utf-8"))
    assert run["input_files"] == ["plan.docx"]
    assert run["metrics"]["llm"]["calls"] == 0
    assert run["metrics"]["analysis_payload_characters"]["total"] > 0
    assert list((output_dir / "document_ir").glob("*.json"))
    assert list((output_dir / "document_summaries").glob("*.json"))
    ir_text = next((output_dir / "document_ir").glob("*.json")).read_text(
        encoding="utf-8"
    )
    assert '"is_merged_copy"' not in ir_text
    assert '"origin_row"' not in ir_text
