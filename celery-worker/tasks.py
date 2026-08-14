import os
import base64
import tempfile
import shutil
from io import BytesIO
import re
import html
from celery import shared_task
from docx import Document
from docx.shared import Pt, RGBColor
from summary_model.report_markup import mark_report_text
from summary_model.web_service import WebPipelineOptions, process_uploaded_documents


REQUIRED_DOCUMENTS = (
    ("plan", "Заявка в план-график"),
    ("contract", "Проект контракта"),
    ("ooz", "ООЗ"),
    ("zapiska", "Пояснительная записка"),
    ("onmck", "ОНМЦК"),
    ("obrasheniye", "Обращение о проведении закупки"),
)
MANDATORY_DOCUMENT_KEYS = {"plan"}


def _add_formatted_runs(paragraph, text: str) -> None:
    tag_pattern = re.compile(r"</?(?:b|u|ins|ok|warn|error|big|doc)>", re.IGNORECASE)
    bold_active = False
    underline_active = False
    ok_active = False
    warn_active = False
    error_active = False
    big_active = False
    doc_active = False
    cursor = 0

    for match in tag_pattern.finditer(text):
        if match.start() > cursor:
            run = paragraph.add_run(html.unescape(text[cursor:match.start()]))
            run.bold = bold_active
            run.underline = underline_active
            if big_active:
                run.font.size = Pt(13)
            if ok_active:
                run.font.color.rgb = RGBColor(0x19, 0x87, 0x54)
            elif warn_active:
                run.font.color.rgb = RGBColor(0xFD, 0x7E, 0x14)
            elif error_active:
                run.font.color.rgb = RGBColor(0xDC, 0x35, 0x45)
            elif doc_active:
                run.font.color.rgb = RGBColor(0x6C, 0x75, 0x7D)

        tag = match.group(0).lower()
        if tag == "<b>":
            bold_active = True
        elif tag == "</b>":
            bold_active = False
        elif tag in ("<u>", "<ins>"):
            underline_active = True
        elif tag in ("</u>", "</ins>"):
            underline_active = False
        elif tag == "<ok>":
            ok_active = True
        elif tag == "</ok>":
            ok_active = False
        elif tag == "<warn>":
            warn_active = True
        elif tag == "</warn>":
            warn_active = False
        elif tag == "<error>":
            error_active = True
        elif tag == "</error>":
            error_active = False
        elif tag == "<big>":
            big_active = True
        elif tag == "</big>":
            big_active = False
        elif tag == "<doc>":
            doc_active = True
        elif tag == "</doc>":
            doc_active = False

        cursor = match.end()

    if cursor < len(text):
        run = paragraph.add_run(html.unescape(text[cursor:]))
        run.bold = bold_active
        run.underline = underline_active
        if big_active:
            run.font.size = Pt(13)
        if ok_active:
            run.font.color.rgb = RGBColor(0x19, 0x87, 0x54)
        elif warn_active:
            run.font.color.rgb = RGBColor(0xFD, 0x7E, 0x14)
        elif error_active:
            run.font.color.rgb = RGBColor(0xDC, 0x35, 0x45)
        elif doc_active:
            run.font.color.rgb = RGBColor(0x6C, 0x75, 0x7D)


def build_result_docx_bytes(ai_response: str) -> bytes:
    """
    Собирает docx-файл из текстового ответа модели с поддержкой таблиц, заголовков и переносов строк.
    """
    document = Document()
    document.add_heading('Результат проверки документов', level=1)

    clean_response = (ai_response or '').replace('\r\n', '\n')
    lines = clean_response.split('\n')

    i = 0
    while i < len(lines):
        line = lines[i].strip()
        if not line:
            i += 1
            continue

        if line.startswith('|') and line.endswith('|'):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith('|') and lines[i].strip().endswith('|'):
                table_lines.append(lines[i].strip())
                i += 1

            rows_data = []
            for tline in table_lines:
                if re.match(r"^\|(?:\s*:?-+:?\s*\|)+$", tline):
                    continue
                cells = [c.strip() for c in tline.strip('|').split('|')]
                rows_data.append(cells)

            if rows_data:
                col_count = max(len(r) for r in rows_data)
                table = document.add_table(rows=len(rows_data), cols=col_count)
                table.style = 'Table Grid'
                for r_idx, row_cells in enumerate(rows_data):
                    for c_idx, cell_value in enumerate(row_cells):
                        if c_idx < col_count:
                            p = table.cell(r_idx, c_idx).paragraphs[0]
                            _add_formatted_runs(p, cell_value)
            continue

        if line.startswith('#### '):
            document.add_heading(line[5:].strip(), level=3)
        elif line.startswith('### '):
            document.add_heading(line[4:].strip(), level=3)
        elif line.startswith('## '):
            document.add_heading(line[3:].strip(), level=2)
        elif line.startswith('# '):
            document.add_heading(line[2:].strip(), level=1)
        elif line == "---":
            p = document.add_paragraph()
            _add_formatted_runs(p, "----------------------------------------")
        else:
            p = document.add_paragraph()
            _add_formatted_runs(p, line)

        i += 1

    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


@shared_task(bind=True, name='rag_worker.process_document_query')
def process_document_query(self, documents):
    try:
        if not isinstance(documents, list):
            raise ValueError("Expected a list of uploaded documents.")

        prepared_documents = []
        present_keys = set()
        for document in documents:
            doc_key = document.get('key')
            doc_label = document.get('label') or doc_key
            file_name = document.get('name')
            file_content_b64 = document.get('content_b64')

            if not doc_key or not file_name or not file_content_b64:
                raise ValueError("Each uploaded document must contain key, name and content_b64.")

            prepared_documents.append({
                'key': doc_key,
                'label': doc_label,
                'name': file_name,
                'content': base64.b64decode(file_content_b64),
            })
            present_keys.add(doc_key)

        missing_docs = [
            label for key, label in REQUIRED_DOCUMENTS
            if key in MANDATORY_DOCUMENT_KEYS and key not in present_keys
        ]
        if missing_docs:
            raise ValueError(
                f"Missing required documents: {', '.join(missing_docs)}."
            )

        temp_dir = tempfile.mkdtemp()
        try:
            pipeline_documents = []
            for index, document in enumerate(prepared_documents, 1):
                key = document['key']
                file_name = os.path.basename(document['name'])
                temp_file_path = os.path.join(temp_dir, f"{index:02d}_{key}_{file_name}")

                with open(temp_file_path, 'wb') as f:
                    f.write(document['content'])

                pipeline_documents.append({
                    'key': key,
                    'label': document['label'],
                    'name': document['name'],
                    'path': temp_file_path,
                })

            pipeline_result = process_uploaded_documents(
                pipeline_documents,
                options=WebPipelineOptions(
                    with_llm_extraction=True,
                    with_semantic_llm=True,
                    with_ktru=True,
                    with_vlm_tables=os.getenv("SUMMARY_WITH_VLM_TABLES", "0") == "1",
                    with_vlm_commercial_offers=os.getenv("SUMMARY_WITH_VLM_COMMERCIAL_OFFERS", "1") == "1",
                    with_vlm_short_documents=os.getenv("SUMMARY_WITH_VLM_SHORT_DOCUMENTS", "1") == "1",
                    ktru_timeout_seconds=int(os.getenv("KTRU_TIMEOUT_SECONDS", "30")),
                    llm_concurrency=int(os.getenv("SUMMARY_LLM_CONCURRENCY", "6")),
                    vlm_max_tables_per_document=int(os.getenv("SUMMARY_VLM_MAX_TABLES_PER_DOCUMENT", "4")),
                    vlm_max_commercial_offer_pages=int(os.getenv("SUMMARY_VLM_MAX_COMMERCIAL_OFFER_PAGES", "8")),
                    vlm_max_short_document_pages=int(os.getenv("SUMMARY_VLM_MAX_SHORT_DOCUMENT_PAGES", "4")),
                ),
            )
            ai_response = mark_report_text(pipeline_result.report_text)
            if pipeline_result.warnings:
                warnings_text = "\n".join(
                    f"- {_public_pipeline_warning(warning)}"
                    for warning in pipeline_result.warnings
                )
                ai_response = f"{ai_response}\n\n<b>Технические предупреждения</b>\n{warnings_text}"
            result_file_bytes = build_result_docx_bytes(ai_response)

            return {
                'ai_response': ai_response,
                'result_file_b64': base64.b64encode(result_file_bytes).decode('utf-8'),
                'result_file_name': 'analysis_result.docx',
                'documents': [
                    {
                        'key': document['key'],
                        'label': document['label'],
                        'name': document['name'],
                    }
                    for document in prepared_documents
                ],
                'status': 'completed'
            }
        finally:
            if os.path.exists(temp_dir):
                shutil.rmtree(temp_dir)
    except Exception as e:
        error_msg = str(e)
        print(f"Error processing documents: {error_msg}")
        raise Exception(error_msg)


def _public_pipeline_warning(warning: str) -> str:
    """Keep report warnings useful without exposing JSON/Pydantic internals."""
    text = " ".join(str(warning or "").split())
    prefix = text.split(":", 1)[0]
    if "спецификация распознана как пустая или шаблонная" in text:
        return f"{prefix}: спецификация не содержит заполненных товарных позиций."
    if "VLM fallback failed" in text:
        localized_prefix = prefix.replace(", table ", ", таблица ")
        return (
            f"{localized_prefix}: сложная таблица не прошла визуальное распознавание; "
            "использован исходный детерминированный разбор."
        )
    if "Structured extraction failed" in text or "structured output (None)" in text:
        return (
            f"{prefix}: языковая модель не вернула структурированный ответ; "
            "использован детерминированный разбор."
        )
    return text[:500] + ("..." if len(text) > 500 else "")
